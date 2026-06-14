"""用户手册 .docx 生成脚本

读取同目录下 用户手册.md,渲染为带页眉页脚的 Word。
运行:python 生成脚本.py
输出:用户手册.docx
"""
from __future__ import annotations
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


HERE = Path(__file__).parent
SRC = HERE / "用户手册.md"
OUT = HERE / "用户手册.docx"

SOFTWARE_NAME = "读享 V1.0 软件"
DOC_VERSION = "V1.0"


def set_run_font(run, size: int = 12) -> None:
    run.font.name = "宋体"
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "宋体")


def add_page_header_footer(section) -> None:
    header = section.header
    h_p = header.paragraphs[0]
    h_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h_run = h_p.add_run(f"{SOFTWARE_NAME} - 用户手册 {DOC_VERSION}")
    set_run_font(h_run, 9)

    footer = section.footer
    f_p = footer.paragraphs[0]
    f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = f_p.add_run("第 ")
    set_run_font(f_run, 9)

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText_page = OxmlElement("w:instrText")
    instrText_page.set(qn("xml:space"), "preserve")
    instrText_page.text = " PAGE "
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    r1 = f_p.add_run()
    r1._element.append(fldChar_begin)
    r1._element.append(instrText_page)
    r1._element.append(fldChar_end)

    of_run = f_p.add_run(" 页 / 共 ")
    set_run_font(of_run, 9)

    fldChar_begin2 = OxmlElement("w:fldChar")
    fldChar_begin2.set(qn("w:fldCharType"), "begin")
    instrText_total = OxmlElement("w:instrText")
    instrText_total.set(qn("xml:space"), "preserve")
    instrText_total.text = " NUMPAGES "
    fldChar_end2 = OxmlElement("w:fldChar")
    fldChar_end2.set(qn("w:fldCharType"), "end")
    r2 = f_p.add_run()
    r2._element.append(fldChar_begin2)
    r2._element.append(instrText_total)
    r2._element.append(fldChar_end2)

    end_run = f_p.add_run(" 页")
    set_run_font(end_run, 9)


def parse_markdown(md_text: str) -> list[tuple[str, str, str]]:
    """返回 [(level, type, content), ...]
    type: 'h1' | 'h2' | 'h3' | 'p' | 'hr' | 'table'
    level: 1/2/3
    """
    blocks: list[tuple[str, str, str]] = []
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("# "):
            blocks.append(("1", "h1", line[2:].strip()))
        elif line.startswith("## "):
            blocks.append(("2", "h2", line[3:].strip()))
        elif line.startswith("### "):
            blocks.append(("3", "h3", line[4:].strip()))
        elif line.startswith("---"):
            blocks.append(("0", "hr", ""))
        elif line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            blocks.append(("0", "table", "\n".join(table_lines)))
            continue
        else:
            para_lines = [line]
            i += 1
            while i < len(lines):
                nxt = lines[i].rstrip()
                if not nxt or nxt.startswith("#") or nxt.startswith("---") or nxt.startswith("|"):
                    break
                para_lines.append(nxt)
                i += 1
            blocks.append(("0", "p", "\n".join(para_lines)))
            continue
        i += 1
    return blocks


def add_table(doc: Document, table_text: str) -> None:
    rows = [r for r in table_text.splitlines() if r.strip() and not re.match(r"^\|[\s\-:|]+\|$", r)]
    if not rows:
        return
    parsed = [re.findall(r"\|([^|]*)", r) for r in rows]
    parsed = [[c.strip() for c in row] for row in parsed]
    n_cols = max(len(row) for row in parsed)
    parsed = [row + [""] * (n_cols - len(row)) for row in parsed]
    table = doc.add_table(rows=len(parsed), cols=n_cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(parsed):
        for c_idx, cell_text in enumerate(row):
            cell_obj = table.cell(r_idx, c_idx)
            cell_obj.text = ""
            p = cell_obj.paragraphs[0]
            run = p.add_run(cell_text)
            set_run_font(run, 10)


def build() -> None:
    md_text = SRC.read_text(encoding="utf-8")
    blocks = parse_markdown(md_text)

    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        add_page_header_footer(section)

    for level, btype, content in blocks:
        if btype == "h1":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content)
            set_run_font(run, 22)
            run.bold = True
        elif btype == "h2":
            p = doc.add_paragraph()
            run = p.add_run(content)
            set_run_font(run, 18)
            run.bold = True
        elif btype == "h3":
            p = doc.add_paragraph()
            run = p.add_run(content)
            set_run_font(run, 14)
            run.bold = True
        elif btype == "p":
            for line in content.splitlines():
                p = doc.add_paragraph()
                run = p.add_run(line)
                set_run_font(run, 12)
        elif btype == "hr":
            doc.add_paragraph("―" * 30)
        elif btype == "table":
            add_table(doc, content)

    doc.save(OUT)
    print(f"生成成功: {OUT}")


if __name__ == "__main__":
    build()