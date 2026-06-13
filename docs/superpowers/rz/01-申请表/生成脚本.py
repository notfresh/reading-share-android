"""软著登记申请表 .docx 生成脚本

读取同目录下 软著登记申请表.md,解析 24 个栏目,生成 Word 文档。
运行:python 生成脚本.py
输出:软著登记申请表.docx
"""
from __future__ import annotations

from datetime import date
from pathlib import Path
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


HERE = Path(__file__).parent
SRC = HERE / "软著登记申请表.md"
OUT = HERE / "软著登记申请表.docx"

# 24 个栏目的标题(在 .md 中以 "## N. " 开头)
SECTION_RE = re.compile(r"^## (\d+)\. (.+)$", re.MULTILINE)


def parse_sections(md_text: str) -> list[tuple[str, str]]:
    """返回 [(栏目号. 标题, 正文), ...] 列表(包含封面 H1 之前丢弃)"""
    sections: list[tuple[str, str]] = []
    matches = list(SECTION_RE.finditer(md_text))
    for i, m in enumerate(matches):
        num_title = f"{m.group(1)}. {m.group(2)}"
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(md_text)
        body = md_text[start:end].strip()
        sections.append((num_title, body))
    return sections


def set_chinese_font(run, size: int = 12) -> None:
    run.font.name = "宋体"
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "宋体")


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_chinese_font(run, 18)
    run.bold = True


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, 14)
    run.bold = True


def add_body(doc: Document, text: str) -> None:
    for line in text.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line.strip() if line else " ")
        set_chinese_font(run, 12)


def build_docx() -> None:
    md_text = SRC.read_text(encoding="utf-8")
    sections = parse_sections(md_text)

    doc = Document()
    # 页面设置:A4
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_title(doc, "计算机软件著作权登记申请表")

    # 23 栏 申请日期 自动填入
    final_sections = []
    for num_title, body in sections:
        if num_title.startswith("23."):
            body = date.today().isoformat()
        final_sections.append((num_title, body))

    for num_title, body in final_sections:
        add_section_heading(doc, num_title)
        add_body(doc, body)

    doc.save(OUT)
    print(f"生成成功: {OUT}")


if __name__ == "__main__":
    build_docx()
